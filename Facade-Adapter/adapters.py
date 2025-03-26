from abc import ABC, abstractmethod 
from pdf2docx import Converter
from docx import Document
import fitz
import subprocess
import os
import sys

# Clase abstracta para los adaptadores
class FileConverter(ABC):
    @abstractmethod
    def convertir(self, origen: str, destino: str):
        pass

# Adaptadores concretos
class PdfToDocxAdapter(FileConverter):
    def convertir(self, origen: str, destino: str):
        cv = Converter(origen)
        cv.convert(destino)
        cv.close()

class DocxToTxtAdapter(FileConverter):
    def convertir(self, origen: str, destino: str):
        doc = Document(origen)
        with open(destino, "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")

class TxtToDocxAdapter(FileConverter):
    def convertir(self, origen: str, destino: str):
        doc = Document()
        with open(origen, "r", encoding="utf-8") as f:
            for line in f:
                doc.add_paragraph(line.strip())
        doc.save(destino)

class PdfToTxtAdapter(FileConverter):
    def convertir(self, origen: str, destino: str):
        doc = fitz.open(origen)
        text = "\n".join([page.get_text() for page in doc])
        with open(destino, "w", encoding="utf-8") as f:
            f.write(text)

class TxtToPdfAdapter(FileConverter):
    def convertir(self, origen: str, destino: str):
        doc = fitz.open()
        text = open(origen, "r", encoding="utf-8").read()
        page = doc.new_page()
        page.insert_text((50, 100), text, fontsize=12)
        doc.save(destino)

class DocxToPdfAdapter(FileConverter):
    def convertir(self, origen: str, destino: str):
        try:
            if sys.platform.startswith("win"):
                # Método para Windows usando LibreOffice
                libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
            else:
                # Método para Linux/macOS
                libreoffice_path = "libreoffice"

            # Ejecutar LibreOffice en modo headless para convertir
            subprocess.run([libreoffice_path, "--headless", "--convert-to", "pdf", origen, "--outdir", os.path.dirname(destino)], check=True)
        except Exception as e:
            raise RuntimeError(f"Error al convertir DOCX a PDF: {str(e)}")